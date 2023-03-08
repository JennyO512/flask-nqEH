from flask import Flask, render_template, request, jsonify
import docx
from docx import Document
from docx2pdf import convert
import os
import openai

app = Flask(__name__)


#@app.route('/')
#def index():
#    return jsonify({"Choo Choo": "Welcome to your Flask app 🚅"})





app = Flask(__name__)


# This route gets the base page, its' like the 1st page that shows up #changed from home to index
@app.route('/', methods=['GET'])
def index():
    return render_template('base2.html')



# write email route 
@app.route('/write-email')
def index():
    return '''
        <form method="POST" action="/search">
            <input type="text" name="query">
            <input type="submit" value="Search">
        </form>
    '''



@app.route('/create-cover-letter', methods=['POST'])
def create_cover_letter():
    # Get the position from the user input
    position = request.form.get('position')
    
    # Get the OpenAI API credentials
    
    openai.api_key = "YOUR API KEY"
    
    # Generate the cover letter using the OpenAI API
    prompt = f"Create a cover letter for {position} position"
    completions = openai.Completion.create(
        engine="text-davinci-002",
        prompt=prompt,
        max_tokens=1024,
        n=1,
        stop=None,
        temperature=0.5,
    )
    letter = completions.choices[0].text
    
    # Create a new word document
    doc = docx.Document()
    
    # Add the cover letter text to the document
    doc.add_paragraph(letter)
    
    # Save the document
    doc.save("Cover_Letter.docx")
    
    
    # Return the cover letter to the textarea in the base2.html template
    return render_template('base2.html', letter=letter, position=position)


@app.route('/create-PDF', methods=['POST'])
def create_PDF():
     # Convert the 'cover_letter.docx' file to PDF
    convert("cover_letter.docx", "cover_letter.pdf")
    
    
    print("cover_letter.docx converted to cover_letter.pdf")




@app.route('/create-resume', methods=['POST'])
def create_resume():
    # Get the position from the user input
    resumecreated = request.form.get('resume')
    
    # Get the OpenAI API credentials
    
    openai.api_key = "your api key"
    
    # Generate the cover letter using the OpenAI API
    prompt = f"Create a resume for {resumecreated}."
    completions = openai.Completion.create(
        engine="text-davinci-002",
        prompt=prompt,
        max_tokens=1024,
        n=1,
        stop=None,
        temperature=0.5,
    )
    resume = completions.choices[0].text
    
    # Create a new word document
    doc = docx.Document()
    
    # Add the resume text to the document
    doc.add_paragraph(resume)
    
    # Save the document
    doc.save('new_resume.docx')
    
    #return 'resume created and saved as resume.docx'
    return render_template('resume.html')


@app.route('/search', methods=['POST'])
def search():
    query = request.form['query']
    openai.api_key = "YOUR API KEY HERE"
    response = openai.Completion.create(engine="text-davinci-002", prompt=f"{query}")
    #return response["choices"][0]["text"]
    return render_template('emails.html', query=query, response=response["choices"][0]["text"])





#if __name__ == '__main__':
#    app.run(debug=True)
    
    
if __name__ == '__main__':
    # Run the Flask app
    app.run(debug=True)

    # Run the Kivy app
   # MyApp().run()
    
#This code uses the Flask web framework to create a simple web application. 
# When the user navigates to the home page, they can enter the position they 
# applied for and submit the form to generate a cover letter. 
# The position is passed to the OpenAI API to generate the cover letter, 
# and the response is then saved to a word document named "cover_letter.docx" in the root of your project.



if __name__ == '__main__':
    app.run(debug=True, port=os.getenv("PORT", default=5000))
